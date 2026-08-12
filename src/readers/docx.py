"""DOCX reader backed by python-docx.

Emits headings, paragraphs, lists, tables, code blocks, and LaTeX-ized OMML
equations.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from ..deps import require
from ..engine.model import CodeBlock, Document, Heading, Paragraph, Quote, Table
from ..math_converters.omml import omath_element_to_latex
from .base import Reader

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_W_R = f"{{{_W_NS}}}r"
_W_T = f"{{{_W_NS}}}t"
_M_OMATH = f"{{{_M_NS}}}oMath"
_M_OMATH_PARA = f"{{{_M_NS}}}oMathPara"


class DocxReader(Reader):
    """Read .docx files into a :class:`Document`."""

    extensions = frozenset({".docx"})
    name = "docx"

    def read(self, path: Path, **kwargs: Any) -> Document:
        require("docx", "reading DOCX files")
        import docx

        doc = docx.Document(path)
        result = Document()

        for element in doc.element.body:
            tag = element.tag
            if tag.endswith("}p"):
                self._parse_paragraph(element, result)
            elif tag.endswith("}tbl"):
                self._parse_table(element, result)

        return result

    def _parse_paragraph(self, p_elem: Any, doc: Document) -> None:
        style_name = ""
        pPr = p_elem.find(f"{{{_W_NS}}}pPr")
        if pPr is not None:
            pStyle = pPr.find(f"{{{_W_NS}}}pStyle")
            if pStyle is not None:
                style_name = (pStyle.get(f"{{{_W_NS}}}val") or "").lower()

        text_parts: list[str] = []
        for child in p_elem:
            tag = child.tag
            if tag == _W_R:
                text_parts.append("".join(t.text or "" for t in child.findall(_W_T)))
            elif tag == _M_OMATH_PARA:
                text_parts.append(f" $${omath_element_to_latex(child)}$$ ")
            elif tag == _M_OMATH:
                text_parts.append(f" ${omath_element_to_latex(child)}$ ")

        full_text = "".join(text_parts).strip()
        if not full_text:
            return

        if style_name.startswith("heading"):
            try:
                level = int(style_name.replace("heading", "").strip())
            except ValueError:
                level = 1
            doc.add_heading(full_text, level=min(level, 6))
        elif "quote" in style_name:
            doc.add_quote(full_text)
        elif "code" in style_name:
            doc.add_code_block(full_text)
        else:
            doc.add_paragraph(full_text)

    def _parse_table(self, tbl_elem: Any, doc: Document) -> None:
        rows: list[list[str]] = []
        for tr in tbl_elem.findall(f"{{{_W_NS}}}tr"):
            row: list[str] = []
            for tc in tr.findall(f"{{{_W_NS}}}tc"):
                cell_text = "".join(tc.itertext()).strip()
                row.append(cell_text)
            if any(row):
                rows.append(row)

        if not rows:
            return

        header = rows[0]
        body = rows[1:] if len(rows) > 1 else []
        doc.add_table(body, header=header)